import pdfplumber
from docx import Document


def _table_to_text(table: list[list[str | None]]) -> str:
    """Convert a pdfplumber table (list of rows) into readable key-value text."""
    lines = []
    for row in table:
        cells = [cell.strip() if cell else "" for cell in row]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def parse_pdf(path: str) -> list[str]:
    """Extract sections from a PDF — one per table, one for free text per page."""
    sections: list[str] = []

    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            tables = page.extract_tables()
            table_bboxes = []
            if tables:
                for tbl_obj in page.find_tables():
                    table_bboxes.append(tbl_obj.bbox)

                for table in tables:
                    text = _table_to_text(table)
                    if text.strip():
                        sections.append(f"[Page {page_num} – Table]\n{text}")

            if table_bboxes:
                remaining_text_parts = []
                full_text = page.extract_text() or ""
                table_texts = set()
                for table in tables:
                    for row in table:
                        for cell in row:
                            if cell and cell.strip():
                                table_texts.add(cell.strip())

                for line in full_text.split("\n"):
                    stripped = line.strip()
                    if stripped and stripped not in table_texts:
                        remaining_text_parts.append(stripped)

                if remaining_text_parts:
                    sections.append(
                        f"[Page {page_num} – Text]\n"
                        + "\n".join(remaining_text_parts)
                    )
            else:
                text = page.extract_text()
                if text and text.strip():
                    sections.append(f"[Page {page_num} – Text]\n{text.strip()}")

    return sections


def parse_docx(path: str) -> list[str]:
    """Extract sections from a DOCX file — paragraphs and tables."""
    sections: list[str] = []
    doc = Document(path)

    # Extract tables
    for i, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        text = "\n".join(rows)
        if text.strip():
            sections.append(f"[Table {i}]\n{text}")

    # Extract paragraphs (group consecutive non-empty paragraphs)
    current_block: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current_block.append(text)
        elif current_block:
            sections.append("[Text]\n" + "\n".join(current_block))
            current_block = []

    if current_block:
        sections.append("[Text]\n" + "\n".join(current_block))

    return sections


def parse_txt(path: str) -> list[str]:
    """Extract sections from a plain text file — split by double newlines."""
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()

    blocks = content.split("\n\n")
    sections = []
    for block in blocks:
        stripped = block.strip()
        if stripped:
            sections.append(f"[Text]\n{stripped}")

    return sections


def parse_document(path: str) -> list[str]:
    """Parse a document based on file extension. Returns list of section strings."""
    ext = path.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return parse_pdf(path)
    elif ext == "docx":
        return parse_docx(path)
    elif ext == "txt":
        return parse_txt(path)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")
